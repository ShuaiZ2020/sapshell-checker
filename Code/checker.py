from docx import Document
import polars as pl
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from glob import glob
from os import path, makedirs
import xlsxwriter
#https://stackoverflow.com/questions/42093013/processing-objects-in-order-in-docx



def get_docx_path(project_path):
    ## list file under prject_path and return first docx file with string of 'tfl'
    pattern = path.join(project_path, 'SapShell/check', '*TFL*.docx')
    files = glob(pattern, recursive=True)
    if files[0]:
        print(f'Found docx file: {files[0]}')
        return files[0]
    else:
        print('No docx file found in the project path.')
        return None



def check_paragraph_font_size(paragraph: Paragraph, font_size =10.5) -> bool:
    if not isinstance(paragraph, Paragraph):
        raise TypeError("Input must be a docx.text.paragraph.Paragraph object")
    # Check explicit paragraph-level font size
    if paragraph.style.font.size and paragraph.style.font.size != Pt(10.5):
        return False

    # Check each run in the paragraph
    for run in paragraph.runs:
        if (run.text!='')|(run is not None) :
            size = run.font.size
        else:
            continue
        if size is not None:
            if size != font_size:
                return False  # Mismatch
    return True



def check_paragraph_font_name(paragraph: Paragraph) -> bool:
    if not isinstance(paragraph, Paragraph):
        raise TypeError("Input must be a docx.text.paragraph.Paragraph object")
    # Check each run's font name
    for run in paragraph.runs:
        run_font = run.font.name
        if (run_font is not None):
            if run_font not in ['宋体', 'Times New Roman']:
                #print(run_font)
                return False

    return True

def is_paragraph_cn_en(paragraph: Paragraph) -> bool:
    if not isinstance(paragraph, Paragraph):
        raise TypeError("Input must be a docx.text.paragraph.Paragraph object")

    full_text = paragraph.text
    parts = full_text.split("^^")

    if len(parts) == 2:
        return True  # Invalid structure (must be exactly one ^^)
    return False

def split_paragraph_cn_en(paragraph: Paragraph) -> bool:
    if not isinstance(paragraph, Paragraph):
        raise TypeError("Input must be a docx.text.paragraph.Paragraph object")

    full_text = paragraph.text
    parts = full_text.split("^^")
    return parts



def get_rows_from_para(paragraph: Paragraph):
    df_dict = {'text':[],
            'style_name':[],
            'good_font_name':[],
            'good_font_size':[],
            'is_cn_en':False,
            'content_type':[],
            'column_id':[],
            'row_id': []
            }
    if not isinstance(paragraph, Paragraph):
        raise TypeError("Input must be a docx.text.paragraph.Paragraph object")
    df_dict['text'].append(paragraph.text)
    df_dict['style_name'].append(paragraph.style.name)
    df_dict['good_font_name'].append(check_paragraph_font_name(paragraph))
    if "toc" in paragraph.style.name:
        df_dict['good_font_size'].append(check_paragraph_font_size(paragraph, 11))
    else:
        df_dict['good_font_size'].append(check_paragraph_font_size(paragraph))
    df_dict['is_cn_en'] = is_paragraph_cn_en(paragraph)
    df_dict['content_type'].append('paragraph')
    df_dict['column_id'].append(None)
    df_dict['row_id'].append(None)
    return(pl.DataFrame(df_dict))
    
def get_df_from_table(table):
    df_dict = {'text':[],
        'style_name':[],
        'good_font_name':[],
        'good_font_size':[],
        'is_cn_en':[],
        'content_type':[],
        'column_id':[],
        'row_id': []
        }
    for row_index in range(len(table.rows)):
        row = table.rows[row_index]
        for cell_index in range(len(row.cells)):
            cell = row.cells[cell_index]
            if (cell.text == '') | (cell is None):
                df_dict['text'].append('')
                df_dict['style_name'].append('')
                df_dict['good_font_name'].append(True)
                df_dict['good_font_size'].append(True)
                df_dict['is_cn_en'].append(False)
                df_dict['content_type'].append('table cell')
                df_dict['column_id'].append(cell_index)
                df_dict['row_id'].append(row_index)
            else:   
                df_dict['text'].append(cell.text)
                df_dict['style_name'].append(cell.paragraphs[0].style.name)
                df_dict['good_font_name'].append(check_paragraph_font_name(cell.paragraphs[0]))
                df_dict['good_font_size'].append(check_paragraph_font_size(cell.paragraphs[0]))
                df_dict['is_cn_en'].append(is_paragraph_cn_en(cell.paragraphs[0]))
                df_dict['content_type'].append('table cell')
                df_dict['column_id'].append(cell_index)
                df_dict['row_id'].append(row_index)
    return pl.DataFrame(df_dict)
    

def get_rows_from_unknowobj(unknowobj):
    df_dict = {'text':[],
        'style_name':[],
        'good_font_name':[],
        'good_font_size':[],
        'is_cn_en':False,
        'content_type':[],
        'column_id':[],
        'row_id': []
        }
    df_dict['text'].append(None)
    df_dict['style_name'].append(None)
    df_dict['good_font_name'].append(None)
    df_dict['good_font_size'].append(None)
    df_dict['is_cn_en'] = False
    df_dict['content_type'].append(None)
    df_dict['column_id'].append(None)
    df_dict['row_id'].append(None)
    return(pl.DataFrame(df_dict))

def print_attributes(obj, include_private=False, include_dunder=False):
    for attr in dir(obj):
        if attr.startswith("__") and attr.endswith("__") and not include_dunder:
            continue
        if attr.startswith("_") and not include_private and not attr.startswith("__"):
            continue
        try:
            value = getattr(obj, attr)
        except Exception as e:
            value = f"<Error: {e}>"
        print(f"{attr}: {value}")

     
def save_df_to_datadir_excel(project_path, df):
    data_path = path.join(project_path, "SapShell/check", 'table_df.xlsx')
    if not path.exists(path.dirname(data_path)):
        makedirs(path.dirname(data_path))

    with xlsxwriter.Workbook(data_path) as wb:
        ws = wb.add_worksheet()

        # Write header
        ws.write_row(0, 0, df.columns)

        # Write rows
        for i, row in enumerate(df.rows()):
            ws.write_row(i+1, 0, row)

    return data_path